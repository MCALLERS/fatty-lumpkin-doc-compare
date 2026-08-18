#!/usr/bin/env python3
"""Edge-case .docx pairs that probe the comparison engine's weak spots."""
import os
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, 'edge')
os.makedirs(OUT, exist_ok=True)


def save(d, name):
    d.save(os.path.join(OUT, name))


def paras(name, lines, style=None):
    d = Document()
    for line in lines:
        d.add_paragraph(line, style=style) if style else d.add_paragraph(line)
    save(d, name)


# 1. short paragraph, one word changed -- bigram similarity is blind here
paras('short_v1.docx', ['Term: three years.', 'Governing law: New York.', 'Logo: old caption text.'])
paras('short_v2.docx', ['Term: five years.', 'Governing law: Delaware.', 'Logo: new caption text.'])

# 2. paragraph split into two
paras('split_v1.docx', ['Alpha beta gamma.', 'Delta epsilon zeta. Eta theta iota.', 'Kappa lambda mu.'])
paras('split_v2.docx', ['Alpha beta gamma.', 'Delta epsilon zeta.', 'Eta theta iota.', 'Kappa lambda mu.'])

# 3. two paragraphs merged into one
paras('merge_v1.docx', ['One two three.', 'Four five six.', 'Seven eight nine.'])
paras('merge_v2.docx', ['One two three.', 'Four five six. Seven eight nine.'])

# 4. numbered list with items deleted from the middle
def numbered(name, items):
    d = Document()
    d.add_heading('Schedule', level=1)
    for it in items:
        d.add_paragraph(it, style='List Number')
    save(d, name)


numbered('list_v1.docx', ['First item.', 'Second item.', 'Third item.', 'Fourth item.', 'Fifth item.'])
numbered('list_v2.docx', ['First item.', 'Third item.', 'Fifth item.'])

# 5. change lives only in a footnote / header / footer
def with_parts(name, footnote_text, header_text, footer_text):
    d = Document()
    d.add_heading('Liability', level=1)
    d.add_paragraph('The cap on liability is set out in the schedule.')
    sec = d.sections[0]
    sec.header.paragraphs[0].text = header_text
    sec.footer.paragraphs[0].text = footer_text
    # python-docx has no footnote API; approximate with a trailing note paragraph
    # in the footer, which exercises the same code path.
    sec.footer.add_paragraph(footnote_text)
    save(d, name)


with_parts('parts_v1.docx', 'Note: cap of $1,000,000.', 'DRAFT v1', 'Confidential - do not circulate')
with_parts('parts_v2.docx', 'Note: cap of $5,000,000.', 'DRAFT v2', 'Confidential - do not circulate')

# 6. heading text changed
def headed(name, heading, body):
    d = Document()
    d.add_heading(heading, level=1)
    d.add_paragraph(body)
    save(d, name)


headed('heading_v1.docx', 'Section 4: Indemnity', 'The indemnity is uncapped.')
headed('heading_v2.docx', 'Section 4: Indemnities', 'The indemnity is uncapped.')

# 7. table with an in-cell edit that must show a margin change bar
def table_doc(name, rate):
    d = Document()
    d.add_heading('Rates', level=1)
    t = d.add_table(rows=0, cols=3)
    t.style = 'Table Grid'
    for row in [('Service', 'Rate', 'Unit'), ('Support', rate, 'per month'), ('Advice', '$900', 'per hour')]:
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.text = v
    save(d, name)


table_doc('table_v1.docx', '$2,400')
table_doc('table_v2.docx', '$3,600')

# 8. one document empty of body text
paras('empty_v1.docx', [''])
paras('empty_v2.docx', ['Now it says something.'])

# 9. long table that must repeat its header across a page break
def long_table(name, changed):
    d = Document()
    d.add_heading('Line items', level=1)
    t = d.add_table(rows=0, cols=3)
    t.style = 'Table Grid'
    cells = t.add_row().cells
    for c, v in zip(cells, ('Item', 'Rate', 'Notes')):
        c.text = v
    for i in range(1, 46):
        cells = t.add_row().cells
        rate = f'${i * 100 + (500 if changed and i % 12 == 0 else 0):,}'
        for c, v in zip(cells, (f'Line item {i}', rate, 'Standard terms apply to this line item.')):
            c.text = v
    save(d, name)


long_table('longtable_v1.docx', False)
long_table('longtable_v2.docx', True)

print('edge fixtures ->', OUT)
