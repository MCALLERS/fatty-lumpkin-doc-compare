#!/usr/bin/env python3
"""Builds realistic .docx test pairs for the compare engine."""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))


def build(path, variant):
    d = Document()
    styles = d.styles
    styles['Normal'].font.name = 'Calibri'
    styles['Normal'].font.size = Pt(11)

    d.add_heading('MASTER SERVICES AGREEMENT', level=0)

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Dated as of ')
    r = p.add_run('12 August 2026' if variant == 'a' else '15 September 2026')
    r.bold = True

    d.add_paragraph(
        'This Master Services Agreement (this "Agreement") is entered into by and between '
        'Bombadillo AI, Inc., a Delaware corporation ("Provider"), and '
        + ('Withywindle Holdings LLC' if variant == 'a' else 'Withywindle Holdings LLC')
        + ' ("Customer"), and is effective as of the date first written above.'
    )

    d.add_heading('1. Services', level=1)
    d.add_paragraph(
        'Provider shall perform the services described in each Statement of Work '
        + ('executed by the parties.' if variant == 'a'
           else 'executed by the parties from time to time, in each case on a non-exclusive basis.')
    )
    d.add_paragraph(
        'Provider shall use commercially reasonable efforts to perform the Services in a '
        'professional and workmanlike manner consistent with '
        + ('industry standards.' if variant == 'a' else 'generally accepted industry standards.')
    )

    d.add_heading('2. Fees and Payment', level=1)
    d.add_paragraph(
        'Customer shall pay Provider the fees set out in the applicable Statement of Work '
        'within ' + ('thirty (30)' if variant == 'a' else 'forty-five (45)') + ' days of receipt of an undisputed invoice.'
    )
    if variant == 'b':
        d.add_paragraph(
            'Any amount not paid when due shall accrue interest at the lesser of one percent '
            '(1%) per month and the maximum rate permitted by applicable law.'
        )

    d.add_heading('3. Term and Termination', level=1)
    for text_a, text_b in [
        ('The initial term of this Agreement shall be twelve (12) months.',
         'The initial term of this Agreement shall be twenty-four (24) months.'),
        ('Either party may terminate this Agreement for material breach upon thirty (30) days written notice.',
         'Either party may terminate this Agreement for material breach upon thirty (30) days written notice, '
         'provided that the breaching party has failed to cure such breach within that period.'),
        ('Sections 4, 5 and 7 shall survive any termination of this Agreement.',
         'Sections 4, 5 and 7 shall survive any termination of this Agreement.'),
    ]:
        d.add_paragraph(text_a if variant == 'a' else text_b, style='List Number')

    if variant == 'a':
        d.add_paragraph(
            'Customer may terminate this Agreement for convenience at any time upon ninety (90) days notice.',
            style='List Number')

    d.add_heading('4. Confidentiality', level=1)
    p = d.add_paragraph()
    p.add_run('Each party ').bold = False
    p.add_run('shall hold in confidence').italic = True
    p.add_run(' all Confidential Information of the other party and shall not disclose it to any third party '
              + ('without prior written consent.' if variant == 'a'
                 else 'without the prior written consent of the disclosing party.'))

    d.add_heading('5. Fee Schedule', level=1)
    rows_a = [
        ('Service', 'Rate', 'Unit'),
        ('Implementation', '$18,000', 'per project'),
        ('Support', '$2,400', 'per month'),
        ('Training', '$1,200', 'per day'),
    ]
    rows_b = [
        ('Service', 'Rate', 'Unit'),
        ('Implementation', '$22,000', 'per project'),
        ('Support', '$2,400', 'per month'),
        ('Advisory', '$3,000', 'per month'),
    ]
    rows = rows_a if variant == 'a' else rows_b
    t = d.add_table(rows=0, cols=3)
    t.style = 'Table Grid'
    for row in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.text = v

    d.add_heading('6. Governing Law', level=1)
    d.add_paragraph(
        'This Agreement shall be governed by the laws of the State of '
        + ('New York' if variant == 'a' else 'Delaware')
        + ', without regard to its conflict of laws principles.'
    )

    if variant == 'b':
        d.add_heading('7. Notices', level=1)
        d.add_paragraph(
            'All notices under this Agreement shall be in writing and delivered by email with '
            'confirmation of receipt, or by nationally recognised overnight courier.'
        )

    d.add_paragraph()
    d.add_paragraph('[Signature page follows]').alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.save(path)


def build_identical(path):
    d = Document()
    d.add_heading('Simple Memo', level=1)
    d.add_paragraph('This memo has not changed at all between versions.')
    d.add_paragraph('Not one word.')
    d.save(path)


def build_big(path, variant):
    d = Document()
    d.add_heading('Schedule of Definitions', level=1)
    for i in range(1, 401):
        txt = f'Section {i}.  The term "Item {i}" means the item identified as number {i} in the register maintained by the Provider.'
        if variant == 'b' and i % 37 == 0:
            txt = txt.replace('maintained by the Provider', 'maintained by the Provider and updated quarterly')
        if variant == 'b' and i % 91 == 0:
            continue
        d.add_paragraph(txt)
    d.save(path)


if __name__ == '__main__':
    build(os.path.join(HERE, 'MSA_v1.docx'), 'a')
    build(os.path.join(HERE, 'MSA_v2.docx'), 'b')
    build_identical(os.path.join(HERE, 'Memo_same_A.docx'))
    build_identical(os.path.join(HERE, 'Memo_same_B.docx'))
    build_big(os.path.join(HERE, 'Definitions_v1.docx'), 'a')
    build_big(os.path.join(HERE, 'Definitions_v2.docx'), 'b')
    print('fixtures written to', HERE)
